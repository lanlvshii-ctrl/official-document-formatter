#!/usr/bin/env python3
"""
ai_structure_analyzer.py
读取 raw.md，调用大语言模型进行：
1. 清理过度分段、修正乱码和 OCR 错误
2. 识别公文结构并打上 HTML 注释格式钩子（<!--HOOK:TYPE-->）
3. 三轮收敛与自检（默认开启），生成勘误建议表
输出 structured.md（+ correction.docx）
"""

import argparse
import json
import os
import re
import sys
import time
from pathlib import Path

try:
    import yaml
except ImportError:
    yaml = None

try:
    import openai
except ImportError:
    openai = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

# ---------------------------------------------------------------------------
# 嵌入提示词
# ---------------------------------------------------------------------------

SYSTEM_PROMPT_FAST = """你是一位专业的中文公文排版专家。你的任务是对输入的原始 Markdown 文本进行两方面的智能处理：

1. **清理与修复**
   - 合并因 PDF 转换或 OCR 导致的过度分段。判断标准：如果一行不以句末标点（。；：！？…）结束，且下一行不是新段落的开头，则应将两行合并为一段。
   - 修正明显的 OCR 乱码和错别字。
   - 保留原文的核心内容和语义，不要增删实质性信息。

2. **结构标注**
   - 为**每一个段落**（包括空行分隔的每个文本块）的开头添加 HTML 注释格式钩子标记。
   - 钩子格式为 `<!--HOOK:类型-->`，必须紧接在段落开头，后面紧跟段落内容，不要加空格。

可用的钩子标记严格限定如下（必须使用这些 exact 字符串）：
- `<!--HOOK:TITLE-->` — 公文的主标题，通常位于文档最开头
- `<!--HOOK:H1-->` — 用汉字编号的大标题，如 "一、总体要求"
- `<!--HOOK:H2-->` — 用括号汉字编号的标题，如 "（一）指导思想"
- `<!--HOOK:H3-->` — 用阿拉伯数字加点的标题，如 "1."
- `<!--HOOK:H4-->` — 用括号阿拉伯数字的标题，如 "（1）"
- `<!--HOOK:BODY-->` — 普通正文段落
- `<!--HOOK:ATTACHMENT-->` — 附件说明，如 "附件：1. xxx"
- `<!--HOOK:SIGNATURE-->` — 发文机关署名
- `<!--HOOK:DATE-->` — 成文日期

### 特殊规则（必须遵守）
- **二级标题拆分**：`<!--HOOK:H2-->` 只包含小标题本身。如果原文中二级标题后面还跟着正文内容（如 "（一）指导思想。坚持以人民为中心..."），必须将其拆分为两个独立段落：第一段是 `<!--HOOK:H2-->（一）指导思想。`，第二段是 `<!--HOOK:BODY-->坚持以人民为中心...`。
- **附件使用正文格式**：`<!--HOOK:ATTACHMENT-->` 的格式与正文一致（仿宋三号），只是内容以"附件："开头。

输出要求：
- 只输出处理后的 Markdown 文本，不要有任何解释、总结、JSON 包装或代码块包裹。
- 每个文本块前都必须有且仅有一个钩子标记。
- 若某段落无法确定类型，默认使用 `<!--HOOK:BODY-->`。
- 保持段落之间的空行，以便后续处理。
"""

SYSTEM_PROMPT_ROUND1 = """你是一位专业的中文公文排版与质检专家。请对输入的原始 Markdown 文本进行以下三方面的深度处理：

### 1. 清理与修复
- 合并因 PDF 转换或 OCR 导致的过度分段。判断标准：如果一行不以句末标点（。；：！？…）结束，且下一行不是新段落的开头，则应将两行合并为一段。
- 修正明显的 OCR 乱码和错别字。
- 保留原文的核心内容和语义，不要增删实质性信息。

### 2. 结构标注
- 为**每一个段落**（空行分隔的每个文本块）添加 HTML 注释格式钩子标记。
- 钩子格式：`<!--HOOK:类型-->`，必须紧接段落开头，后面**不加空格**直接跟内容。
- 可用钩子（exact 字符串）：
  `<!--HOOK:TITLE-->`、`<!--HOOK:H1-->`、`<!--HOOK:H2-->`、`<!--HOOK:H3-->`、`<!--HOOK:H4-->`、`<!--HOOK:BODY-->`、`<!--HOOK:ATTACHMENT-->`、`<!--HOOK:SIGNATURE-->`、`<!--HOOK:DATE-->`
- 若无法确定类型，默认 `<!--HOOK:BODY-->`。

### 特殊规则（必须遵守）
- **二级标题拆分**：`<!--HOOK:H2-->` 只包含小标题本身。如果原文中二级标题后面还跟着正文内容（如 "（一）指导思想。坚持以人民为中心..."），必须将其拆分为两个独立段落：第一段是 `<!--HOOK:H2-->（一）指导思想。`，第二段是 `<!--HOOK:BODY-->坚持以人民为中心...`。
- **附件使用正文格式**：`<!--HOOK:ATTACHMENT-->` 的格式与正文一致（仿宋三号），只是内容以"附件："开头。

### 3. 质量勘误
- 逐段检查文本，发现以下类型的问题并记录：
  错别字、句法错误、逻辑错误、数字错误、时间错误、立场错误、术语错误、政治错误、标注不当
- 对每处问题记录：段落位置（从1开始计数）、原文片段、问题类型、问题说明、修改建议。

### 输出格式（严格遵守）
直接输出处理后的 Markdown 文本（每段带钩子，段间空一行）。
然后另起一行输出 JSON 代码块，格式如下：
```json
{
  "corrections": [
    {
      "paragraph": 1,
      "original": "原文片段",
      "type": "错别字",
      "description": "具体说明",
      "suggestion": "修改建议"
    }
  ]
}
```
不要输出"第一部分"、"第二部分"等提示词本身。
"""

SYSTEM_PROMPT_ROUND2 = """你是一位公文格式规范化专家。请对以下带钩子的文本进行收敛处理：

1. 检查每段的钩子标注是否准确，修正任何不当标注（如把正文标成标题、把附件标成正文等）。
2. 确保每个文本块都有且仅有一个合法钩子。
3. 若某段无钩子，在其开头补 `<!--HOOK:BODY-->`。
4. 删除钩子与内容之间的任何空格。
5. 保持段落结构和原文内容不变，不增删实质性信息。
6. **二级标题必须拆分**：若 `<!--HOOK:H2-->` 段落中包含小标题后的正文内容（如 "（一）指导思想。坚持以人民为中心..."），必须拆分为两段：第一段保留 `<!--HOOK:H2-->（一）指导思想。`，第二段改为 `<!--HOOK:BODY-->坚持以人民为中心...`。
7. **附件使用正文格式**：`<!--HOOK:ATTACHMENT-->` 仅做内容标记，不应用黑体等特殊字体样式。

合法钩子清单：`<!--HOOK:TITLE-->`、`<!--HOOK:H1-->`、`<!--HOOK:H2-->`、`<!--HOOK:H3-->`、`<!--HOOK:H4-->`、`<!--HOOK:BODY-->`、`<!--HOOK:ATTACHMENT-->`、`<!--HOOK:SIGNATURE-->`、`<!--HOOK:DATE-->`

只输出处理后的带钩子 Markdown 文本，不要有任何解释、总结或 JSON。
"""

SYSTEM_PROMPT_ROUND3 = """你是一位公文质检专家。请对以下带钩子的文本进行最终自检：

1. **标注核对**：逐段确认钩子标注是否正确（TITLE/H1/H2/H3/H4/BODY/ATTACHMENT/SIGNATURE/DATE）。
2. **二级标题拆分检查**：确保所有 `<!--HOOK:H2-->` 只包含小标题本身，正文内容已拆分为独立的 `<!--HOOK:BODY-->`。
3. **内容质检**：再次检查内容中的以下问题：
   错别字、句法错误、逻辑错误、数字错误、时间错误、立场错误、术语错误、政治错误。
4. **整理输出**：输出最终处理结果和勘误建议表。

### 输出格式（严格遵守）
直接输出最终 Markdown 文本（每段带钩子，段间空一行）。
然后另起一行输出 JSON 代码块，格式如下（如无问题则 corrections 为空数组）：
```json
{
  "corrections": [
    {
      "paragraph": 3,
      "original": "原文片段",
      "type": "数字错误",
      "description": "具体说明",
      "suggestion": "修改建议"
    }
  ]
}
```
不要输出"第一部分"、"第二部分"等提示词本身。
"""

# ---------------------------------------------------------------------------
# 配置管理
# ---------------------------------------------------------------------------

def ensure_config_dir() -> Path:
    """确保配置目录存在（使用用户级目录，避免敏感信息存入项目仓库）"""
    config_dir = Path.home() / ".config" / "official-document-formatter"
    config_dir.mkdir(parents=True, exist_ok=True)
    return config_dir


def load_api_config() -> dict:
    """加载 API 配置，返回字典（可能为空）"""
    if yaml is None:
        print("⚠️  缺少 pyyaml，无法读取/保存配置文件，请执行: pip install pyyaml")
        return {}
    config_file = ensure_config_dir() / "api_config.yaml"
    if config_file.exists():
        try:
            return yaml.safe_load(config_file.read_text(encoding="utf-8")) or {}
        except Exception as e:
            print(f"⚠️  读取配置文件失败: {e}，将使用交互式输入")
            return {}
    return {}


def save_api_config(config: dict) -> None:
    """保存 API 配置到文件"""
    if yaml is None:
        print("⚠️  缺少 pyyaml，无法保存配置文件")
        return
    config_file = ensure_config_dir() / "api_config.yaml"
    try:
        with config_file.open("w", encoding="utf-8") as f:
            yaml.dump(config, f, allow_unicode=True, sort_keys=False)
    except Exception as e:
        print(f"⚠️  保存配置文件失败: {e}")


def get_api_credentials() -> dict | None:
    """
    获取 API 凭证。优先级：
    1. 环境变量
    2. ~/.config/official-document-formatter/api_config.yaml
    3. 常见宿主环境默认变量（OPENAI_API_KEY / ANTHROPIC_API_KEY / DEEPSEEK_API_KEY）
    若均未找到，返回 None（由调用方决定：独立模式报错，宿主模式由宿主 AI 接管）
    返回 {"api_key": str, "base_url": str, "model": str, "model_reasoning": str|None}
    """
    api_key = os.environ.get("OFFICIAL_FORMATTER_API_KEY", "").strip()
    base_url = os.environ.get("OFFICIAL_FORMATTER_BASE_URL", "").strip()
    model = os.environ.get("OFFICIAL_FORMATTER_MODEL", "").strip()
    model_reasoning = os.environ.get("OFFICIAL_FORMATTER_MODEL_REASONING", "").strip() or None

    if api_key:
        print("ℹ️  已从环境变量读取 API 配置")
        return {
            "api_key": api_key,
            "base_url": base_url or "https://api.openai.com/v1",
            "model": model or "gpt-4o",
            "model_reasoning": model_reasoning,
        }

    config = load_api_config()
    api_key = config.get("api_key", "").strip()
    if api_key:
        print("ℹ️  已从配置文件读取 API 配置")
        return {
            "api_key": api_key,
            "base_url": config.get("base_url", "https://api.openai.com/v1"),
            "model": config.get("model", "gpt-4o"),
            "model_reasoning": config.get("model_reasoning") or None,
        }

    # 尝试常见宿主环境默认变量
    for env_key, default_base, default_model in [
        ("OPENAI_API_KEY", "https://api.openai.com/v1", "gpt-4o"),
        ("ANTHROPIC_API_KEY", "https://api.anthropic.com/v1", "claude-3-5-sonnet-20241022"),
        ("DEEPSEEK_API_KEY", "https://api.deepseek.com/v1", "deepseek-chat"),
    ]:
        api_key = os.environ.get(env_key, "").strip()
        if api_key:
            print(f"ℹ️  已从宿主环境变量 {env_key} 读取 API 配置")
            return {
                "api_key": api_key,
                "base_url": default_base,
                "model": default_model,
                "model_reasoning": None,
            }

    # 无可用外部 API，返回 None，由调用方处理
    return None


def _prompt_save_credentials(api_key: str, base_url: str, model: str, model_reasoning: str | None) -> None:
    """（保留给命令行工具调用，技能本身不再主动询问）"""
    config_path = ensure_config_dir() / "api_config.yaml"
    print(f"ℹ️  提示：可通过环境变量或 {config_path} 持久化配置，避免重复输入。")


# ---------------------------------------------------------------------------
# LLM 调用
# ---------------------------------------------------------------------------

def estimate_tokens(text: str) -> int:
    cn_chars = len(re.findall(r"[\u4e00-\u9fa5]", text))
    other_chars = len(text) - cn_chars
    return int(cn_chars * 1.5 + other_chars * 0.25)


def truncate_for_context(text: str, max_tokens: int = 60000) -> str:
    est = estimate_tokens(text)
    if est <= max_tokens:
        return text
    ratio = max_tokens / est
    keep_chars = int(len(text) * ratio * 0.95)
    text = text[:keep_chars]
    # 优先按段落边界截断，其次按行截断
    last_break = text.rfind("\n\n")
    if last_break > len(text) * 0.5:
        text = text[:last_break]
    else:
        line_break = text.rfind("\n")
        if line_break > len(text) * 0.5:
            text = text[:line_break]
    print(f"⚠️  文档过长（约 {est} tokens），已按段落边界截断至约 {max_tokens} tokens")
    return text


def strip_think_tags(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"<analysis>.*?</analysis>", "", text, flags=re.DOTALL)
    return text


def call_llm(raw_text: str, credentials: dict, system_prompt: str, max_retries: int = 3, model_override: str = None) -> str:
    if openai is None:
        raise RuntimeError("缺少 openai 库，请执行: pip install openai")

    raw_text = truncate_for_context(raw_text)
    client = openai.OpenAI(api_key=credentials["api_key"], base_url=credentials["base_url"])
    model = model_override or credentials["model"]

    last_error = None
    for attempt in range(1, max_retries + 1):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": raw_text},
                ],
                temperature=0.2,
            )
            return strip_think_tags(response.choices[0].message.content.strip())
        except openai.RateLimitError as e:
            last_error = e
            wait = 2 ** attempt
            print(f"⚠️  速率限制，{wait} 秒后重试（第 {attempt}/{max_retries} 次）...")
            time.sleep(wait)
        except openai.APIConnectionError as e:
            last_error = e
            if attempt < max_retries:
                wait = 2 ** attempt
                print(f"⚠️  网络连接失败，{wait} 秒后重试（第 {attempt}/{max_retries} 次）...")
                time.sleep(wait)
            else:
                raise
        except Exception as e:
            last_error = e
            raise

    raise RuntimeError(f"LLM 调用在 {max_retries} 次重试后仍失败: {last_error}")


# ---------------------------------------------------------------------------
# 结果解析
# ---------------------------------------------------------------------------

def parse_markdown_and_corrections(text: str) -> tuple[str, list]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    json_match = re.search(r"```json\s*(\{.*?\})\s*```", text, re.DOTALL)
    if not json_match:
        json_match = re.search(r"```\s*(\{.*?\})\s*```", text, re.DOTALL)

    corrections = []
    if json_match:
        json_str = json_match.group(1)
        try:
            data = json.loads(json_str)
            corrections = data.get("corrections", []) if isinstance(data, dict) else []
        except json.JSONDecodeError:
            pass
        markdown_text = (text[:json_match.start()] + text[json_match.end():]).strip()
    else:
        markdown_text = text.strip()

    return markdown_text, corrections


def post_process(text: str) -> str:
    text = re.sub(r"^```[\w]*\n?", "", text)
    text = re.sub(r"\n?```$", "", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 确保每个 HOOK 前面有空行分隔（处理 LLM 输出的连续单行钩子）
    text = re.sub(r'\n(<!--HOOK:\w+-->)', r'\n\n\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 去掉所有数字编号后的空格（如 "1. " → "1."）
    text = re.sub(r'(?<!\d)(\d+)\.(?!\d)\s+', r'\1.', text)

    valid_hooks = {
        "<!--HOOK:TITLE-->", "<!--HOOK:H1-->", "<!--HOOK:H2-->",
        "<!--HOOK:H3-->", "<!--HOOK:H4-->", "<!--HOOK:BODY-->",
        "<!--HOOK:ATTACHMENT-->", "<!--HOOK:SIGNATURE-->", "<!--HOOK:DATE-->",
    }

    paragraphs = text.split("\n\n")
    fixed = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        has_hook = any(para.startswith(hook) for hook in valid_hooks)
        if not has_hook:
            para = "<!--HOOK:BODY-->" + para
        else:
            para = re.sub(r"^(<!--HOOK:\w+-->)\s+", r"\1", para)
        fixed.append(para)

    return "\n\n".join(fixed)


# ---------------------------------------------------------------------------
# 勘误表生成
# ---------------------------------------------------------------------------

def _make_border_element(tag, val="single", sz="6", space="0", color="000000"):
    """构造一个 w:XXX 边框子元素"""
    el = OxmlElement(tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    return el


def _set_table_borders(table):
    """
    给表格设置边框：外框粗线（sz=12），内部细线（sz=6）
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        tblBorders.append(_make_border_element(f"w:{side}", sz="12"))
    for side in ("insideH", "insideV"):
        tblBorders.append(_make_border_element(f"w:{side}", sz="6"))
    tblPr.append(tblBorders)


def _set_cell_fixed_width(cell, width_emu: int):
    """显式设置单元格宽度（dxa 单位）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    old_cw = tcPr.find(qn("w:tcW"))
    if old_cw is not None:
        tcPr.remove(old_cw)
    tcW = OxmlElement("w:tcW")
    # 1 EMU = 1/635 twip (dxa)
    twips_val = int(width_emu / 635)
    tcW.set(qn("w:w"), str(twips_val))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def generate_correction_docx(corrections: list, output_path: Path, doc_title: str = "勘误建议表"):
    if not HAS_DOCX:
        print("⚠️  缺少 python-docx，跳过生成勘误建议表。请执行: pip install python-docx")
        return

    doc = Document()

    # 设置为横向A4（参考 word_table_landscape.py）
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(doc_title)
    title_run.font.name = "微软雅黑"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if not corrections:
        info = doc.add_paragraph()
        info_run = info.add_run("经自检，暂未发现明显问题。")
        info_run.font.name = "微软雅黑"
        info_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        info_run.font.size = Pt(11)
        info_run.italic = True
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.save(str(output_path))
        return

    headers = ["序号", "段落位置", "原文内容", "问题类型", "问题说明", "修改建议"]
    table_data = [headers]
    for idx, item in enumerate(corrections, start=1):
        para_num = str(item.get("paragraph", ""))
        original = str(item.get("original", ""))
        q_type = str(item.get("type", ""))
        desc = str(item.get("description", ""))
        suggestion = str(item.get("suggestion", ""))
        table_data.append([str(idx), para_num, original, q_type, desc, suggestion])

    table = doc.add_table(rows=len(table_data), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    # 横向A4可用宽度 ≈ 297 - 25*2 = 247mm，均分给6列
    usable_width_mm = 247
    num_cols = len(headers)
    col_width_mm = usable_width_mm / num_cols
    col_width_emu = Mm(col_width_mm)

    _set_table_borders(table)

    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # 清空默认内容后重新添加，以便统一设置字体
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            _set_cell_fixed_width(cell, col_width_emu)

        if i == 0:
            for cell in row.cells:
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "2C5282")
                cell._tc.get_or_add_tcPr().append(shading_elm)

        row.height = Pt(30)

    doc.add_paragraph()
    stats_para = doc.add_paragraph()
    stats_text = f"共发现 {len(corrections)} 处问题，请结合实际语境判断是否需要修改。"
    stats_run = stats_para.add_run(stats_text)
    stats_run.font.name = "微软雅黑"
    stats_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    stats_run.font.size = Pt(9)
    stats_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    stats_run.italic = True

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------

def run_fast_mode(raw_text: str, credentials: dict) -> str:
    print("🤖 正在调用模型进行快速结构分析...")
    result = call_llm(raw_text, credentials, SYSTEM_PROMPT_FAST)
    return post_process(result)


def run_thorough_mode(raw_text: str, credentials: dict) -> tuple[str, list]:
    model_r1 = credentials.get("model_reasoning") or credentials["model"]
    if not credentials.get("model_reasoning"):
        print(f"ℹ️  未配置独立 reasoning 模型，Round 1 将使用默认模型: {model_r1}")

    print(f"🤖 Round 1/3 深度标注与勘误（模型: {model_r1}）...")
    r1_output = call_llm(raw_text, credentials, SYSTEM_PROMPT_ROUND1, model_override=model_r1)
    r1_md, r1_corrections = parse_markdown_and_corrections(r1_output)
    r1_md = post_process(r1_md)
    print(f"   Round 1 完成，初步发现 {len(r1_corrections)} 处问题")

    print("🤖 Round 2/3 收敛规范化...")
    r2_output = call_llm(r1_md, credentials, SYSTEM_PROMPT_ROUND2)
    r2_md = post_process(r2_output)
    print("   Round 2 完成")

    print("🤖 Round 3/3 最终自检与勘误...")
    r3_output = call_llm(r2_md, credentials, SYSTEM_PROMPT_ROUND3)
    r3_md, r3_corrections = parse_markdown_and_corrections(r3_output)
    final_md = post_process(r3_md)
    print(f"   Round 3 完成，最终勘误 {len(r3_corrections)} 处问题")

    seen = set()
    merged = []
    for c in r1_corrections + r3_corrections:
        key = (c.get("paragraph"), c.get("original"), c.get("type"))
        if key not in seen:
            seen.add(key)
            merged.append(c)

    if merged:
        print(f"📋 合并去重后共 {len(merged)} 处勘误建议")
    return final_md, merged


def main():
    parser = argparse.ArgumentParser(description="AI 结构分析器：raw.md → structured.md（+ correction.docx）")
    parser.add_argument("input", help="输入 raw.md 路径")
    parser.add_argument("-o", "--output", help="输出 structured.md 路径（默认与输入同目录）")
    parser.add_argument("--api-key", help="直接指定 API Key（覆盖其他来源）")
    parser.add_argument("--base-url", help="直接指定 API Base URL")
    parser.add_argument("--model", help="直接指定模型名称")
    parser.add_argument("--model-reasoning", help="直接指定深度思考模型名称（用于 Round 1）")
    parser.add_argument("--fast", action="store_true", help="使用单轮快速模式（不生成勘误表）")
    parser.add_argument("--no-correction", action="store_true", help="即使使用精细模式，也不生成 correction.docx")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)

    raw_text = input_path.read_text(encoding="utf-8")
    if not raw_text.strip():
        print("❌ raw.md 为空文件")
        sys.exit(1)

    if args.api_key:
        credentials = {
            "api_key": args.api_key.strip(),
            "base_url": (args.base_url or "https://api.openai.com/v1"),
            "model": (args.model or "gpt-4o"),
            "model_reasoning": args.model_reasoning or None,
        }
    else:
        credentials = get_api_credentials()
        if args.base_url:
            credentials["base_url"] = args.base_url
        if args.model:
            credentials["model"] = args.model
        if args.model_reasoning:
            credentials["model_reasoning"] = args.model_reasoning

    if args.fast:
        structured_text = run_fast_mode(raw_text, credentials)
        corrections = []
    else:
        structured_text, corrections = run_thorough_mode(raw_text, credentials)

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = input_path.parent / f"{input_path.stem.replace('_raw', '')}_structured.md"
        if output_path.name == input_path.name:
            output_path = input_path.parent / f"{input_path.stem}_structured.md"

    output_path.write_text(structured_text, encoding="utf-8")
    print(f"✅ 标注完成: {output_path}")

    if not args.fast and not args.no_correction and corrections is not None:
        correction_path = output_path.parent / "勘误建议表.docx"
        generate_correction_docx(corrections, correction_path, doc_title="公文勘误建议表")
        if HAS_DOCX:
            print(f"✅ 勘误表生成: {correction_path}")

    hook_counts = {}
    for line in structured_text.splitlines():
        m = re.match(r"^(<!--HOOK:\w+-->)", line)
        if m:
            hook_counts[m.group(1)] = hook_counts.get(m.group(1), 0) + 1
    if hook_counts:
        print("   标注统计:")
        for hook, count in sorted(hook_counts.items()):
            print(f"      {hook}: {count} 段")


if __name__ == "__main__":
    main()
