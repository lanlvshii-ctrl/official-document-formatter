#!/usr/bin/env python3
"""
docx_formatter.py
读取带 HTML 注释钩子的 structured.md，直接生成严格符合国标格式的 Word 文档。
无需 VBA，Python 直接完成全部排版。
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

# ---------------------------------------------------------------------------
# 字体检测
# ---------------------------------------------------------------------------

def _scan_system_fonts() -> set:
    """扫描常见系统字体目录，返回已安装的字体名集合（不区分大小写，去掉扩展名）"""
    candidates = set()
    dirs = []
    import platform
    system = platform.system()
    home = Path.home()
    if system == "Darwin":
        dirs = [
            home / "Library/Fonts",
            Path("/Library/Fonts"),
            Path("/System/Library/Fonts"),
            Path("/System/Library/Fonts/Cache"),
        ]
    elif system == "Windows":
        dirs = [Path("C:/Windows/Fonts")]
    else:
        dirs = [home / ".fonts", Path("/usr/share/fonts"), Path("/usr/local/share/fonts")]

    for d in dirs:
        if not d.exists():
            continue
        for f in d.rglob("*"):
            if f.suffix.lower() in (".ttf", ".ttc", ".otf", ".dfont"):
                candidates.add(f.stem)
    return candidates


_SYSTEM_FONTS = None

def get_system_fonts() -> set:
    global _SYSTEM_FONTS
    if _SYSTEM_FONTS is None:
        _SYSTEM_FONTS = _scan_system_fonts()
    return _SYSTEM_FONTS


def resolve_font(primary: str, *fallbacks: str) -> str:
    """如果 primary 字体未安装，依次尝试 fallback，否则返回 primary"""
    installed = get_system_fonts()
    for name in (primary,) + fallbacks:
        if name in installed:
            return name
    return primary


# ---------------------------------------------------------------------------
# OOXML 辅助
# ---------------------------------------------------------------------------

def _set_doc_grid(section):
    """设置文档网格：每行 28 字，每页 22 行（标准 OOXML 值）"""
    sectPr = section._sectPr
    docGrid = sectPr.find(qn("w:docGrid"))
    if docGrid is None:
        docGrid = OxmlElement("w:docGrid")
        sectPr.append(docGrid)
    docGrid.set(qn("w:type"), "linesAndChars")
    docGrid.set(qn("w:linePitch"), "312")
    docGrid.set(qn("w:charSpace"), "-1042")


def _add_page_number_field(paragraph):
    """在段落中插入 Word 页码域（PAGE）"""
    run = paragraph.add_run("- ")
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)
    run = paragraph.add_run(" -")


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False):
    """设置 Run 字体，同时设置中西文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _delete_paragraph(paragraph):
    """从文档中删除指定段落"""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


# ---------------------------------------------------------------------------
# 解析 structured.md
# ---------------------------------------------------------------------------

def parse_structured_md(text: str) -> list[dict]:
    """
    解析带 <!--HOOK:TYPE--> 注释的 markdown。
    返回 [{"hook": "TITLE", "content": "..."}, ...]
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 兼容 LLM 输出的连续单行钩子（没有空行分隔的情况）
    text = re.sub(r'\n(<!--HOOK:\w+-->)', r'\n\n\1', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    valid_hooks = {
        "TITLE", "H1", "H2", "H3", "H4",
        "BODY", "ATTACHMENT", "SIGNATURE", "DATE",
    }
    paragraphs = []

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        m = re.match(r"^<!--HOOK:(\w+)-->", block)
        if m and m.group(1) in valid_hooks:
            hook = m.group(1)
            content = block[len(f"<!--HOOK:{hook}-->"):].lstrip()
            paragraphs.append({"hook": hook, "content": content})
        else:
            paragraphs.append({"hook": "BODY", "content": block})

    return paragraphs


def validate_hook_mapping(paragraphs: list[dict]) -> list[dict]:
    """
    根据段落内容的编号格式，自动检测并修正错误的钩子标注。
    
    规则（来自国标公文编号规范）：
    - "一、" "二、" → 必须是 H1
    - "（一）" "（二）" → 必须是 H2
    - "1．" "2．" → 必须是 H3
    - "（1）" "（2）" → 必须是 H4
    
    如果 AI 标错了（如把"一、"标成 H2），此函数会自动修正并打印警告。
    """
    # 编号格式 → 正确钩子的映射
    NUMBERING_RULES = [
        (r'^[一二三四五六七八九十]+、', "H1"),   # 一、二、三、
        (r'^（[一二三四五六七八九十]+）', "H2"),  # （一）（二）（三）
        (r'^\d+．', "H3"),                       # 1．2．3．
        (r'^（\d+）', "H4"),                     # （1）（2）（3）
    ]
    
    corrected = []
    corrections = []
    
    for para in paragraphs:
        content = para["content"].strip()
        hook = para["hook"]
        expected_hook = None
        
        for pattern, correct_hook in NUMBERING_RULES:
            if re.match(pattern, content):
                expected_hook = correct_hook
                break
        
        if expected_hook and hook != expected_hook and hook in ("H1", "H2", "H3", "H4"):
            corrections.append(f"  ⚠️ 修正: <!--HOOK:{hook}-->{content[:20]}… → <!--HOOK:{expected_hook}-->")
            corrected.append({"hook": expected_hook, "content": para["content"]})
        else:
            corrected.append(para)
    
    if corrections:
        print(f"🔧 钩子校验：检测到 {len(corrections)} 处标注错误，已自动修正：")
        for c in corrections:
            print(c)
    else:
        print("✅ 钩子校验：所有标题钩子标注正确")
    
    return corrected


def split_heading_body(paragraphs: list[dict]) -> list[dict]:
    """
    对 H1/H2 段落进行兜底拆分：若内容中包含小标题后的正文，拆分为标题 + BODY。
    H1 例：一、总体要求。为深入贯彻落实... → H1 + BODY
    H2 例：（一）指导思想。坚持以人民为中心... → H2 + BODY
    """
    result = []
    for para in paragraphs:
        if para["hook"] in ("H1", "H2"):
            content = para["content"]
            if para["hook"] == "H1":
                # H1 匹配：一、标题。正文 / 一、标题：正文
                m = re.match(r'^([一二三四五六七八九十]+、[^。；：！？…]*[。；：！？…])(.+)$', content)
            else:
                # H2 匹配：（一）标题。正文 / （一）标题：正文
                m = re.match(r'^(（[一二三四五六七八九十]+）[^。；：！？…]*[。；：！？…])(.+)$', content)
            if m:
                result.append({"hook": para["hook"], "content": m.group(1)})
                result.append({"hook": "BODY", "content": m.group(2).strip()})
            else:
                result.append(para)
        else:
            result.append(para)
    return result


def split_paragraphs_by_newline(paragraphs: list[dict]) -> list[dict]:
    """
    将包含 \n 的段落拆分为多个段落，确保 Word 中所有换行都是段落标记 ^p，
    而不是手动换行符 ^l。
    例外：Markdown 表格块保持完整，不拆分。
    """
    result = []
    for para in paragraphs:
        if _looks_like_markdown_table(para["content"]):
            result.append(para)
            continue
        lines = para["content"].split("\n")
        for i, line in enumerate(lines):
            line = line.strip()
            if not line and i == len(lines) - 1:
                continue
            result.append({"hook": para["hook"], "content": line})
    return result


# ---------------------------------------------------------------------------
# Markdown 表格 → DOCX 表格
# ---------------------------------------------------------------------------

def _looks_like_markdown_table(text: str) -> bool:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) < 2:
        return False
    return "|" in lines[0] and "|" in lines[1] and "---" in lines[1]


def _add_markdown_table(doc: Document, text: str, body_font: str):
    """将 Markdown 表格文本转换为 docx 表格，应用仿宋小四（12pt）国标格式。"""
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if len(lines) < 2:
        return

    rows = []
    for line in lines:
        if line.startswith("|") and line.endswith("|"):
            line = line[1:-1]
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)

    # 第二行是分隔符，跳过
    data_rows = [rows[0]] + rows[2:]
    max_cols = max((len(r) for r in data_rows), default=0)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(data_rows), cols=max_cols)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    for i, row_data in enumerate(data_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= max_cols:
                break
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            _set_run_font(run, body_font, 12)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    # 尝试让表格宽度占满页面（约 16cm）
    try:
        table_width = Cm(16)
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(table_width)))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)
    except Exception:
        pass


# ---------------------------------------------------------------------------
# 格式化核心
# ---------------------------------------------------------------------------

def _add_formatted_paragraph(
    doc: Document,
    content: str,
    font_name: str,
    size_pt: int,
    alignment,
    line_spacing_pt: int,
    first_line_indent_cm: float,
    bold: bool = False,
    keep_with_next: bool = False,
):
    """统一添加一个格式化的段落"""
    p = doc.add_paragraph()
    run = p.add_run(content)
    _set_run_font(run, font_name, size_pt, bold=bold)
    p.alignment = alignment
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(line_spacing_pt)
    p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    return p


def apply_format(doc: Document, paragraphs: list[dict]):
    """将国标格式应用到 docx 文档"""
    body_font = resolve_font("仿宋", "FangSong")
    title_font = resolve_font("方正小标宋简体", "黑体", "SimHei")
    h1_font = resolve_font("黑体", "SimHei", "微软雅黑")
    h2_font = resolve_font("楷体", "KaiTi")
    page_num_font = resolve_font("宋体", "SimSun", "Songti SC")

    # 校验钩子标注（根据编号格式自动修正错误钩子）
    paragraphs = validate_hook_mapping(paragraphs)
    # 兜底拆分 H1/H2，并拆分内部换行，避免 ^l
    paragraphs = split_heading_body(paragraphs)
    paragraphs = split_paragraphs_by_newline(paragraphs)

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(3.0)
    section.footer_distance = Cm(2.8)

    # 文档网格
    _set_doc_grid(section)

    # 页脚页码（居中，宋体四号，"- {num} -"）
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    _add_page_number_field(footer_para)
    for run in footer_para.runs:
        _set_run_font(run, page_num_font, 14)

    # 逐段格式化
    for idx, para_info in enumerate(paragraphs):
        hook = para_info["hook"]
        content = para_info["content"]

        if hook == "TITLE":
            _add_formatted_paragraph(
                doc, content, title_font, 22, WD_ALIGN_PARAGRAPH.CENTER, 32, 0
            )
        elif hook == "H1":
            _add_formatted_paragraph(
                doc, content, h1_font, 16, WD_ALIGN_PARAGRAPH.JUSTIFY, 28, 0.74, bold=False
            )
        elif hook == "H2":
            _add_formatted_paragraph(
                doc, content, h2_font, 16, WD_ALIGN_PARAGRAPH.JUSTIFY, 28, 0.74, bold=False
            )
        elif hook == "H3":
            _add_formatted_paragraph(
                doc, content, body_font, 16, WD_ALIGN_PARAGRAPH.JUSTIFY, 28, 0.74, bold=True
            )
        elif hook == "H4":
            _add_formatted_paragraph(
                doc, content, body_font, 16, WD_ALIGN_PARAGRAPH.JUSTIFY, 28, 0.74, bold=True
            )
        elif hook == "BODY" or hook == "ATTACHMENT":
            if _looks_like_markdown_table(content):
                _add_markdown_table(doc, content, body_font)
            else:
                _add_formatted_paragraph(
                    doc, content, body_font, 16, WD_ALIGN_PARAGRAPH.JUSTIFY, 28, 0.74
                )
        elif hook == "SIGNATURE":
            p = _add_formatted_paragraph(
                doc, content, body_font, 16, WD_ALIGN_PARAGRAPH.RIGHT, 28, 0
            )
            if idx + 1 < len(paragraphs) and paragraphs[idx + 1]["hook"] == "DATE":
                p.paragraph_format.keep_with_next = True
        elif hook == "DATE":
            _add_formatted_paragraph(
                doc, content, body_font, 16, WD_ALIGN_PARAGRAPH.RIGHT, 28, 0
            )

    # 删除末尾空白段落（避免产生空白页）
    while doc.paragraphs:
        last = doc.paragraphs[-1]
        if last.text.strip():
            break
        _delete_paragraph(last)


# ---------------------------------------------------------------------------
# 主入口
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="structured.md → final.docx（Python 直接国标排版）")
    parser.add_argument("input", help="输入 structured.md 路径")
    parser.add_argument("-o", "--output", help="输出 final.docx 路径（默认与输入同目录）")
    args = parser.parse_args()

    if not HAS_DOCX:
        print("❌ 缺少 python-docx，请执行: pip install python-docx")
        sys.exit(1)

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)

    text = input_path.read_text(encoding="utf-8")
    paragraphs = parse_structured_md(text)

    if not paragraphs:
        print("❌ structured.md 中没有解析到任何段落")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = input_path.parent / f"公文排版后-{input_path.stem}.docx"

    doc = Document()
    apply_format(doc, paragraphs)
    doc.save(str(output_path))
    print(f"✅ 排版完成: {output_path}")
    print(f"   共 {len(paragraphs)} 个段落")


if __name__ == "__main__":
    main()
